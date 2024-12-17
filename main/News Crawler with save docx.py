import requests
from bs4 import BeautifulSoup
from googletrans import Translator
import tkinter as tk
import customtkinter as ctk
import webbrowser
from docx import Document  # Word 파일 작성을 위한 라이브러리
from tkinter import filedialog  # 파일 저장 경로를 선택하기 위한 모듈

translator = Translator()

def fetch_recent_articles(base_url):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36'
    }
    response = requests.get(base_url, headers=headers)
    
    articles = []
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')

        for article_block in soup.find_all('div', class_='news-analysis-v2_content__z0iLP'):
            try:
                title_tag = article_block.find('a', {'data-test': 'article-title-link'})
                title = title_tag.text.strip()
                link = title_tag['href']
                full_link = f"https://www.investing.com{link}" if not link.startswith('https') else link

                description_tag = article_block.find('p', {'data-test': 'article-description'})
                description = description_tag.text.strip() if description_tag else "요약 없음"

                author_tag = article_block.find('span', {'data-test': 'news-provider-name'})
                time_tag = article_block.find('time', {'data-test': 'article-publish-date'})
                author = author_tag.text.strip() if author_tag else "작성자 없음"
                upload_time = time_tag['datetime'] if time_tag else "시간 정보 없음"

                articles.append({
                    'title': title,
                    'link': full_link,
                    'summary': description,
                    'author': author,
                    'upload_time': upload_time
                })

                if len(articles) >= 20:
                    break

            except Exception as e:
                print(f"오류 발생: {e}")
                continue
    return articles

class NewsDashboard(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("뉴스 요약 대시보드")
        self.geometry("1000x800")
        ctk.set_appearance_mode("light")

        self.news_data = {
            "Stock Market News": [],
            "Economic Indicators": [],
            "Economy News": []
        }
        self.setup_ui()
        self.load_articles()

    def setup_ui(self):
        self.title_label = ctk.CTkLabel(self, text="📰 최근 뉴스 대시보드 (최대 20개 기사)", font=("Arial", 28, "bold"))
        self.title_label.pack(pady=15)

        self.tabview = ctk.CTkTabview(self, corner_radius=15)
        self.tabview.pack(expand=True, fill="both", padx=15, pady=10)

        self.tabs = {}
        for category in self.news_data.keys():
            self.tabs[category] = self.tabview.add(category)
            self.setup_tab(self.tabs[category])

        self.summary_frame = ctk.CTkFrame(self)
        self.summary_frame.pack(expand=True, fill="x", padx=15, pady=10)

        self.summary_textbox = ctk.CTkTextbox(self.summary_frame, font=("Arial", 16), height=150, wrap="word")
        self.summary_textbox.pack(expand=True, fill="both", padx=10, pady=5)

        self.button_frame = ctk.CTkFrame(self.summary_frame)
        self.button_frame.pack(pady=10)

        self.link_button = ctk.CTkButton(self.button_frame, text="🔗 원본 기사 보기", command=self.open_article_link, font=("Arial", 14))
        self.link_button.pack(side="left", padx=10)

        self.refresh_button = ctk.CTkButton(self.button_frame, text="🔄 새로고침", command=self.refresh_news, font=("Arial", 14))
        self.refresh_button.pack(side="left", padx=10)

        self.save_button = ctk.CTkButton(self.button_frame, text="💾 문서 저장", command=self.save_to_word, font=("Arial", 14))
        self.save_button.pack(side="left", padx=10)

    def setup_tab(self, tab):
        listbox = tk.Listbox(tab, font=("Arial", 16), height=20, selectmode='multiple', exportselection=False)
        listbox.pack(expand=True, fill="both", padx=10, pady=10)
        listbox.bind("<<ListboxSelect>>", lambda event, lb=listbox: self.show_article_details(event, lb))
        tab.listbox = listbox

    def load_articles(self):
        self.news_data["Stock Market News"] = fetch_recent_articles('https://www.investing.com/news/stock-market-news')
        self.news_data["Economic Indicators"] = fetch_recent_articles('https://www.investing.com/news/economic-indicators')
        self.news_data["Economy News"] = fetch_recent_articles('https://www.investing.com/news/economy')

        for category, articles in self.news_data.items():
            listbox = self.tabs[category].listbox
            listbox.delete(0, 'end')
            for article in articles:
                listbox.insert('end', article['title'])

    def refresh_news(self):
        self.refresh_button.configure(state="disabled", text="🔄 로딩 중...")
        self.summary_textbox.delete("1.0", "end")
        self.link_button.configure(state="disabled")
        self.after(100, self.load_articles)
        self.after(2000, lambda: self.refresh_button.configure(state="normal", text="🔄 새로고침"))

    def show_article_details(self, event, listbox):
        selection = listbox.curselection()
        if selection:
            idx = selection[0]
            category = self.tabview.get()
            article = self.news_data[category][idx]
            english_summary = article['summary']
            korean_summary = translator.translate(english_summary, src='en', dest='ko').text

            self.summary_textbox.delete("1.0", "end")
            self.summary_textbox.insert("1.0", f"🗒️ 영어 요약:\n{english_summary}\n\n🌐 한글 번역:\n{korean_summary}")
            self.current_article_link = article['link']
            self.link_button.configure(state="normal")

    def open_article_link(self):
        if hasattr(self, 'current_article_link'):
            webbrowser.open(self.current_article_link)

    def save_to_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word 파일", "*.docx")])
        if not file_path:
            return  # 사용자가 취소한 경우

        doc = Document()
        doc.add_heading("선택된 뉴스 요약 문서", level=1)

        current_tab = self.tabview.get()
        listbox = self.tabs[current_tab].listbox
        selection = listbox.curselection()

        for idx in selection:
            article = self.news_data[current_tab][idx]
            korean_summary = translator.translate(article['summary'], src='en', dest='ko').text
            doc.add_paragraph(f"제목: {article['title']}")
            doc.add_paragraph(f"작성자: {article['author']}")
            doc.add_paragraph(f"🔹 영어 요약: {article['summary']}")
            doc.add_paragraph(f"🔹 한글 번역: {korean_summary}")
            doc.add_paragraph(f"링크: {article['link']}")
            doc.add_paragraph("-" * 50)

        doc.save(file_path)
        print(f"문서가 '{file_path}'로 저장되었습니다.")

if __name__ == "__main__":
    app = NewsDashboard()
    app.mainloop()